"""
document_processor.py
======================
Extracts plain text from uploaded PDF, DOCX, and TXT files.

Strategy:
    PDF  -> try pdfplumber first (best layout/text fidelity).
            If a page returns no text (i.e. it's a scanned image), fall back
            to OCR (pytesseract) on that page's rasterized image.
            If pdfplumber itself fails to open the file, fall back to pypdf.
    DOCX -> python-docx, paragraphs + tables.
    TXT  -> decode with a few common encodings, fallback to errors='ignore'.

All public functions return a ProcessingResult so callers (app.py) always
get a consistent shape back, whether extraction succeeded, partially
succeeded (some OCR pages empty), or failed with a validation error.
"""

import io
import os
from dataclasses import dataclass, field
from typing import Optional

MAX_FILE_SIZE_MB = 15
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}


@dataclass
class ProcessingResult:
    success: bool
    text: str = ""
    file_type: str = ""
    pages_processed: int = 0
    ocr_pages_used: int = 0
    warnings: list = field(default_factory=list)
    error: Optional[str] = None


class DocumentValidationError(Exception):
    pass


def validate_file(filename: str, file_bytes: bytes):
    """Raise DocumentValidationError if the file is not acceptable."""
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise DocumentValidationError(
            f"Unsupported file type '{ext}'. Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
        )
    size_mb = len(file_bytes) / (1024 * 1024)
    if size_mb > MAX_FILE_SIZE_MB:
        raise DocumentValidationError(f"File is {size_mb:.1f}MB; max allowed is {MAX_FILE_SIZE_MB}MB.")
    if len(file_bytes) == 0:
        raise DocumentValidationError("File is empty.")
    return ext


def process_document(filename: str, file_bytes: bytes) -> ProcessingResult:
    """Main entry point: validates, routes to the right extractor, returns ProcessingResult."""
    try:
        ext = validate_file(filename, file_bytes)
    except DocumentValidationError as e:
        return ProcessingResult(success=False, error=str(e))

    try:
        if ext == ".pdf":
            return _process_pdf(file_bytes)
        elif ext == ".docx":
            return _process_docx(file_bytes)
        elif ext == ".txt":
            return _process_txt(file_bytes)
    except Exception as e:  # noqa: BLE001 - want a graceful, user-facing message
        return ProcessingResult(success=False, error=f"Failed to process document: {e}")

    return ProcessingResult(success=False, error="Unhandled file type.")


# ------------------------------------------------------------------ PDF ----
def _process_pdf(file_bytes: bytes) -> ProcessingResult:
    warnings = []
    ocr_pages_used = 0
    pages_text = []

    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages_text.append(text)
                else:
                    # Likely a scanned/image-only page -> OCR fallback
                    ocr_text = _ocr_pdf_page(page)
                    if ocr_text.strip():
                        pages_text.append(ocr_text)
                        ocr_pages_used += 1
                    else:
                        warnings.append(f"Page {i + 1}: no extractable text (image quality too low for OCR).")
        full_text = "\n\n".join(pages_text)
        if not full_text.strip():
            return ProcessingResult(
                success=False,
                error="Could not extract any text from this PDF, even with OCR. It may be corrupted or blank.",
            )
        return ProcessingResult(
            success=True, text=full_text, file_type="pdf",
            pages_processed=len(pages_text), ocr_pages_used=ocr_pages_used, warnings=warnings,
        )
    except Exception as plumber_err:
        # pdfplumber failed to even open the file -> try pypdf as a fallback engine
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            pages_text = [p.extract_text() or "" for p in reader.pages]
            full_text = "\n\n".join(pages_text)
            if not full_text.strip():
                return ProcessingResult(success=False, error="PDF appears to contain no extractable text.")
            warnings.append(f"Used fallback PDF engine (pdfplumber error: {plumber_err}).")
            return ProcessingResult(
                success=True, text=full_text, file_type="pdf",
                pages_processed=len(pages_text), ocr_pages_used=0, warnings=warnings,
            )
        except Exception as pypdf_err:
            return ProcessingResult(
                success=False,
                error=f"Both PDF engines failed. pdfplumber: {plumber_err} | pypdf: {pypdf_err}",
            )


def _ocr_pdf_page(page) -> str:
    """Rasterize a pdfplumber page to an image and run pytesseract on it."""
    try:
        import pytesseract
        im = page.to_image(resolution=200).original
        return pytesseract.image_to_string(im)
    except Exception:
        return ""


# ----------------------------------------------------------------- DOCX ----
def _process_docx(file_bytes: bytes) -> ProcessingResult:
    import docx

    document = docx.Document(io.BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    full_text = "\n".join(parts)
    if not full_text.strip():
        return ProcessingResult(success=False, error="DOCX file appears to be empty.")
    return ProcessingResult(success=True, text=full_text, file_type="docx", pages_processed=1)


# ------------------------------------------------------------------ TXT ----
def _process_txt(file_bytes: bytes) -> ProcessingResult:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            text = file_bytes.decode(encoding)
            break
        except (UnicodeDecodeError, LookupError):
            continue
    else:
        text = file_bytes.decode("utf-8", errors="ignore")

    if not text.strip():
        return ProcessingResult(success=False, error="Text file is empty.")
    return ProcessingResult(success=True, text=text, file_type="txt", pages_processed=1)
